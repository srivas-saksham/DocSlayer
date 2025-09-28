from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from pygments import highlight
from pygments.lexers import get_lexer_for_filename, PythonLexer
from pygments.formatter import Formatter
from pygments.token import Token
import os
import json


# -------------------------------
# Color Scheme Mapping
# -------------------------------

COLOR_MAP = {
    Token.Keyword: RGBColor(0, 0, 255),          # Blue - keywords (def, class, if, etc.)
    Token.Keyword.Constant: RGBColor(0, 0, 255), # Blue - True, False, None
    Token.Keyword.Namespace: RGBColor(0, 0, 255), # Blue - import, from
    Token.Name.Function: RGBColor(128, 0, 128),   # Purple - function names
    Token.Name.Class: RGBColor(128, 0, 128),      # Purple - class names
    Token.Name.Builtin: RGBColor(0, 128, 128),    # Teal - built-in functions
    Token.String: RGBColor(0, 128, 0),            # Green - strings
    Token.String.Doc: RGBColor(0, 128, 0),        # Green - docstrings
    Token.Comment: RGBColor(128, 128, 128),       # Gray - comments
    Token.Comment.Single: RGBColor(128, 128, 128), # Gray - single line comments
    Token.Comment.Multiline: RGBColor(128, 128, 128), # Gray - multiline comments
    Token.Number: RGBColor(255, 140, 0),          # Orange - numbers
    Token.Number.Integer: RGBColor(255, 140, 0),  # Orange - integers
    Token.Number.Float: RGBColor(255, 140, 0),    # Orange - floats
    Token.Operator: RGBColor(0, 0, 0),            # Black - operators (+, -, =, etc.)
    Token.Punctuation: RGBColor(0, 0, 0),         # Black - punctuation
    Token.Name.Variable: RGBColor(0, 0, 0),       # Black - variables
    Token.Name.Attribute: RGBColor(148, 0, 211),  # Dark Violet - attributes
    Token.Literal: RGBColor(255, 20, 147),        # Deep Pink - literals
    Token.Error: RGBColor(255, 0, 0),             # Red - errors
}

FALLBACK_COLOR = RGBColor(0, 0, 0)  # Black for unrecognized tokens


# -------------------------------
# Helpers
# -------------------------------

def apply_shading(cell, color="D9D9D9"):
    """Apply background shading to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def apply_cell_borders(cell):
    """Apply outside borders only to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    # Define border style for all four sides
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # Border width
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')  # Black border
        tcBorders.append(border)
    
    tcPr.append(tcBorders)


def set_doc_margins(doc, margin_inch=0.5):
    """Set page margins for all sections."""
    for section in doc.sections:
        section.top_margin = Inches(margin_inch)
        section.bottom_margin = Inches(margin_inch)
        section.left_margin = Inches(margin_inch)
        section.right_margin = Inches(margin_inch)


def apply_font(run, size=11, bold=False, font_name="Consolas"):
    """Apply universal font style."""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # default black


def get_lexer_from_file(file_path):
    """Get appropriate lexer based on file extension with fallback to Python."""
    try:
        lexer = get_lexer_for_filename(file_path, stripall=True)
        return lexer
    except Exception:
        # Fallback to Python lexer if detection fails
        return PythonLexer(stripall=True)


# -------------------------------
# Enhanced Syntax Highlighting
# -------------------------------

class DocxFormatter(Formatter):
    """Custom formatter to apply consistent syntax highlighting into DOCX runs."""
    def __init__(self, paragraph, font_name="Consolas", font_size=11):
        super().__init__()
        self.paragraph = paragraph
        self.font_name = font_name
        self.font_size = font_size

    def format(self, tokensource, outfile):
        for ttype, value in tokensource:
            run = self.paragraph.add_run(value)
            run.font.name = self.font_name
            run.font.size = Pt(self.font_size)
            
            # Apply color based on token type with exact matching
            color = self._get_token_color(ttype)
            run.font.color.rgb = color
            
            # Apply italic for comments
            if ttype in Token.Comment:
                run.italic = True
            
            # Apply bold for keywords and function definitions
            if ttype in Token.Keyword or ttype in Token.Name.Function:
                run.bold = True

    def _get_token_color(self, ttype):
        """Get color for token type with hierarchical matching."""
        # Try exact match first
        if ttype in COLOR_MAP:
            return COLOR_MAP[ttype]
        
        # Try parent token types (e.g., Token.String.Single -> Token.String)
        current = ttype
        while current.parent:
            if current in COLOR_MAP:
                return COLOR_MAP[current]
            current = current.parent
        
        # Check if token is in any of the main categories
        for token_type, color in COLOR_MAP.items():
            if ttype in token_type:
                return color
        
        # Fallback to black
        return FALLBACK_COLOR


def add_syntax_highlighted_code(doc, content, file_path, font_size=11):
    """Add syntax highlighted code into a DOCX paragraph with consistent coloring."""
    code_para = doc.add_paragraph()
    
    # Get lexer based on file extension
    lexer = get_lexer_from_file(file_path)
    
    formatter = DocxFormatter(code_para, font_size=font_size)
    highlight(content, lexer, formatter)
    return code_para


def add_syntax_highlighted_code_to_cell(cell, content, file_path, font_size=11):
    """Add syntax highlighted code into a table cell with consistent coloring."""
    # Clear existing content
    cell.paragraphs[0].clear()
    
    # Get lexer based on file extension
    lexer = get_lexer_from_file(file_path)
    
    formatter = DocxFormatter(cell.paragraphs[0], font_size=font_size)
    highlight(content, lexer, formatter)


# -------------------------------
# New Helper Functions for Global Options
# -------------------------------

def add_credentials(doc, credentials):
    """Add student details section at the beginning of document."""
    # Title
    title = doc.add_paragraph("Student Details")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    apply_font(run, size=16, bold=True)
    
    # Create table for credentials
    table = doc.add_table(rows=len(credentials), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, (key, value) in enumerate(credentials.items()):
        # Format key nicely (capitalize and replace underscores)
        formatted_key = key.replace('_', ' ').title() + ":"
        
        # Add key and value
        key_cell = table.rows[i].cells[0]
        value_cell = table.rows[i].cells[1]
        
        # Apply borders to both cells
        apply_cell_borders(key_cell)
        apply_cell_borders(value_cell)
        
        # Set alignment for key cell (center both horizontally and vertically)
        key_paragraph = key_cell.paragraphs[0]
        key_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        key_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Set alignment for value cell (center horizontally, top vertically)
        value_paragraph = value_cell.paragraphs[0]
        value_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        value_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        
        key_run = key_paragraph.add_run(formatted_key)
        apply_font(key_run, size=12, bold=True)
        
        value_run = value_paragraph.add_run(value)
        apply_font(value_run, size=12)
        
        # Apply shading to key cells
        apply_shading(key_cell, "E8E8E8")
    
    # Add some space after credentials
    doc.add_paragraph()


def add_index_page(doc, index_fields, num_questions):
    """Add index page with table format."""
    # Title
    title = doc.add_paragraph("Index")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    apply_font(run, size=16, bold=True)
    
    # Convert frontend index_fields format to field list
    if isinstance(index_fields, str):
        try:
            index_fields = json.loads(index_fields)
        except:
            index_fields = {}
    
    # Map frontend boolean fields to display names
    field_mapping = {
        'sno': 'S. No.',
        'topic': 'Topic', 
        'date': 'Date',
        'teacherSignature': "Teacher's Signature"
    }
    
    # Build field list from enabled fields
    enabled_fields = []
    if isinstance(index_fields, dict):
        for key, enabled in index_fields.items():
            if enabled and key in field_mapping:
                enabled_fields.append(field_mapping[key])
    
    # Fallback if no fields enabled
    if not enabled_fields:
        enabled_fields = ["S. No.", "Topic", "Date", "Teacher's Signature"]
    
    # Create table with headers + rows for each question
    table = doc.add_table(rows=num_questions + 1, cols=len(enabled_fields))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Add headers
    for i, field in enumerate(enabled_fields):
        cell = table.rows[0].cells[i]
        apply_cell_borders(cell)  # Add borders to header cells
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Horizontal center
        run = paragraph.add_run(field)
        apply_font(run, size=12, bold=True)
        apply_shading(cell, "D9D9D9")
        
        # Set vertical alignment for header cells
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Add rows with auto-filled data
    for row_idx in range(1, num_questions + 1):
        for col_idx, field in enumerate(enabled_fields):
            cell = table.rows[row_idx].cells[col_idx]
            apply_cell_borders(cell)  # Add borders to data cells
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Horizontal center for data cells
            
            # Set vertical alignment for data cells (top aligned)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            
            # Auto-fill S. No. and Topic columns
            if field.lower() in ["s. no.", "s.no.", "serial no.", "sr. no."]:
                run = paragraph.add_run(str(row_idx))
                apply_font(run, size=11)
            elif field.lower() in ["topic", "question", "problem"]:
                run = paragraph.add_run(f"Question {row_idx}")
                apply_font(run, size=11)
            # Leave Date and Teacher's Signature empty for manual filling
    
    # Add page break after index
    doc.add_page_break()

def add_page_numbering(doc):
    """Add page numbering to all sections."""
    for section in doc.sections:
        # Access the footer
        footer = section.footer
        
        # Clear existing footer content
        footer.paragraphs[0].clear()
        
        # Add page numbering
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page number field
        run = footer_para.add_run("Page ")
        apply_font(run, size=10)
        
        # Add page number field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        run = footer_para.add_run(" of ")
        apply_font(run, size=10)
        
        # Add total pages field
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        
        instrText2 = OxmlElement('w:instrText')
        instrText2.text = "NUMPAGES"
        
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar3)
        run._r.append(instrText2)
        run._r.append(fldChar4)


# -------------------------------
# Templates with Border Support
# -------------------------------

# Example of how to modify apply_template1 in docx_generator.py

def apply_template1(doc, idx, content, file_path, highlight=False, ai_output=None):
    """Template 1: Question heading, shaded 'Source Code', code block, then output if available."""
    if idx > 1:
        doc.add_page_break()

    # Heading (centered)
    heading = doc.add_paragraph(f"Question {idx}")
    heading.alignment = 1  # center
    run = heading.runs[0]
    apply_font(run, size=16, bold=True)

    # Source Code Section
    # Subheading in shaded cell
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    apply_cell_borders(cell)
    run = cell.paragraphs[0].add_run("Source Code")
    apply_font(run, size=14, bold=True)
    apply_shading(cell, "D9D9D9")

    # Code block
    code_table = doc.add_table(rows=1, cols=1)
    code_cell = code_table.rows[0].cells[0]
    apply_cell_borders(code_cell)

    if highlight:
        add_syntax_highlighted_code_to_cell(code_cell, content, file_path, font_size=12)
    else:
        run = code_cell.paragraphs[0].add_run(content)
        apply_font(run, size=12)

    # AI Output Section (NEW - only if ai_output is provided)
    if ai_output:
        # Add some spacing between source and output
        doc.add_paragraph()
        
        # Output subheading in shaded cell
        output_table = doc.add_table(rows=1, cols=1)
        output_cell = output_table.rows[0].cells[0]
        apply_cell_borders(output_cell)
        run = output_cell.paragraphs[0].add_run("Output")
        apply_font(run, size=14, bold=True)
        apply_shading(output_cell, "D9D9D9")

        # Output content block
        output_content_table = doc.add_table(rows=1, cols=1)
        output_content_cell = output_content_table.rows[0].cells[0]
        apply_cell_borders(output_content_cell)

        # Add the AI output as plain text (it's just a string)
        run = output_content_cell.paragraphs[0].add_run(ai_output)
        apply_font(run, size=11, font_name="Consolas")


def apply_template2(doc, idx, content, file_path, highlight=False, ai_output=None):
    """Template 2: Left-aligned heading with one empty line, Source Code label with character shading, code in table."""
    if idx > 1:
        doc.add_page_break()

    # Heading: "Question {idx}" - Bold, default size (~14px), left-aligned
    heading = doc.add_paragraph(f"Question {idx}")
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.runs[0]
    apply_font(run, size=14, bold=True)
    
    # Add exactly one empty line after heading
    doc.add_paragraph()
    
    # Source Code Label with character shading (not table shading)
    source_label_para = doc.add_paragraph()
    source_run = source_label_para.add_run("Source Code")
    apply_font(source_run, size=16, bold=False)  # 16px, not bold
    
    # Apply character shading using OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "D9D9D9")  # Light gray background
    source_run._r.get_or_add_rPr().append(shd)
    
    # Code block in 1x1 table - no extra blank lines
    code_table = doc.add_table(rows=1, cols=1)
    code_cell = code_table.rows[0].cells[0]
    apply_cell_borders(code_cell)  # Add borders
    
    if highlight:
        add_syntax_highlighted_code_to_cell(code_cell, content, file_path, font_size=12)
    else:
        # Clear the default paragraph and add code with proper formatting
        code_cell.paragraphs[0].clear()
        run = code_cell.paragraphs[0].add_run(content)
        apply_font(run, size=12, bold=False, font_name="Consolas")

    # AI Output Section (NEW - only if ai_output is provided)
    if ai_output:
        # Add exactly one empty line before output
        doc.add_paragraph()
        
        # Output Label with character shading (matching source code style)
        output_label_para = doc.add_paragraph()
        output_run = output_label_para.add_run("Output")
        apply_font(output_run, size=16, bold=False)  # 16px, not bold
        
        # Apply character shading using OxmlElement
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9D9D9")  # Light gray background
        output_run._r.get_or_add_rPr().append(shd)
        
        # Output content block in 1x1 table
        output_content_table = doc.add_table(rows=1, cols=1)
        output_content_cell = output_content_table.rows[0].cells[0]
        apply_cell_borders(output_content_cell)
        
        # Clear the default paragraph and add output with proper formatting
        output_content_cell.paragraphs[0].clear()
        run = output_content_cell.paragraphs[0].add_run(ai_output)
        apply_font(run, size=12, bold=False, font_name="Consolas")


def apply_template3(doc, idx, content, file_path, highlight=False, ai_output=None):
    """Template 3: Minimalist style (Problem + code)."""
    if idx > 1:
        doc.add_page_break()

    heading = doc.add_paragraph(f"Problem {idx}")
    run = heading.runs[0]
    apply_font(run, size=12, bold=True)

    if highlight:
        add_syntax_highlighted_code(doc, content, file_path, font_size=11)
    else:
        code_para = doc.add_paragraph()
        run = code_para.add_run(content)
        apply_font(run, size=11)

    # AI Output Section (NEW - only if ai_output is provided)
    if ai_output:
        # Output heading (minimalist style)
        output_heading = doc.add_paragraph("Output")
        run = output_heading.runs[0]
        apply_font(run, size=12, bold=True)
        
        # Output content (plain paragraph, minimalist)
        output_para = doc.add_paragraph()
        run = output_para.add_run(ai_output)
        apply_font(run, size=11, font_name="Consolas")


def apply_template4(doc, idx, content, file_path, highlight=False, ai_output=None):
    """Template 4: Academic/lab report style with formal structure."""
    if idx > 1:
        doc.add_page_break()

    # Heading (left-aligned, Times New Roman, 14pt, bold, underlined)
    heading = doc.add_paragraph(f"Question {idx}")
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.runs[0]
    apply_font(run, size=14, bold=True, font_name="Times New Roman")
    run.underline = True

    # Add some spacing after heading
    doc.add_paragraph()

    # Aim section
    aim_para = doc.add_paragraph()
    aim_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    aim_label = aim_para.add_run("Aim: ")
    apply_font(aim_label, size=12, bold=True, font_name="Times New Roman")
    aim_text = aim_para.add_run("Write aim here.")
    apply_font(aim_text, size=12, font_name="Times New Roman")

    # Algorithm section
    algo_para = doc.add_paragraph()
    algo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    algo_label = algo_para.add_run("Algorithm: ")
    apply_font(algo_label, size=12, bold=True, font_name="Times New Roman")
    algo_text = algo_para.add_run("Steps of solution here.")
    apply_font(algo_text, size=12, font_name="Times New Roman")

    # Source Code subheading in shaded cell
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    apply_cell_borders(cell)  # Add borders
    run = cell.paragraphs[0].add_run("Source Code")
    apply_font(run, size=14, bold=True, font_name="Times New Roman")
    apply_shading(cell, "D9D9D9")

    # Code block in table cell
    code_table = doc.add_table(rows=1, cols=1)
    code_cell = code_table.rows[0].cells[0]
    apply_cell_borders(code_cell)  # Add borders

    if highlight:
        add_syntax_highlighted_code_to_cell(code_cell, content, file_path, font_size=10)
    else:
        run = code_cell.paragraphs[0].add_run(content)
        apply_font(run, size=10, font_name="Consolas")

    # AI Output Section (NEW - only if ai_output is provided)
    if ai_output:
        # Add some spacing between source and output
        doc.add_paragraph()
        
        # Output subheading in shaded cell (academic style)
        output_table = doc.add_table(rows=1, cols=1)
        output_cell = output_table.rows[0].cells[0]
        apply_cell_borders(output_cell)
        run = output_cell.paragraphs[0].add_run("Output")
        apply_font(run, size=14, bold=True, font_name="Times New Roman")
        apply_shading(output_cell, "D9D9D9")

        # Output content block in table cell
        output_content_table = doc.add_table(rows=1, cols=1)
        output_content_cell = output_content_table.rows[0].cells[0]
        apply_cell_borders(output_content_cell)

        # Add the AI output with academic formatting
        run = output_content_cell.paragraphs[0].add_run(ai_output)
        apply_font(run, size=10, font_name="Consolas")

    # Add horizontal line (simulated with thin paragraph spacing)
    doc.add_paragraph()
    
    # Teacher's signature placeholder
    signature_para = doc.add_paragraph()
    signature_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_run = signature_para.add_run("Teacher's Signature: ____________________")
    apply_font(sig_run, size=11, font_name="Times New Roman")


def apply_template5(doc, idx, content, file_path, highlight=False, ai_output=None):
    """Template 5: Creative colorful style with fun presentation elements."""
    if idx > 1:
        doc.add_page_break()

    # Decorative border around entire question (outer table)
    border_table = doc.add_table(rows=1, cols=1)
    border_cell = border_table.rows[0].cells[0]
    apply_cell_borders(border_cell)  # Add borders
    
    # Alternate background shading for each question block
    if idx % 2 == 0:
        apply_shading(border_cell, "F0F8FF")  # Alice Blue for even questions
    else:
        apply_shading(border_cell, "F5F5F5")  # White Smoke for odd questions
    
    # Clear the border cell and add content inside
    border_cell.paragraphs[0].clear()
    
    # Heading with emoji (centered, colored)
    heading_para = border_cell.add_paragraph(f"🎨 Creative Challenge Q{idx}")
    heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading_para.runs[0]
    apply_font(run, size=18, bold=True, font_name="Segoe UI")
    run.font.color.rgb = RGBColor(0, 128, 128)  # Teal color

    # Add extra spacing
    border_cell.add_paragraph()

    # Subheading in shaded cell (sky blue background)
    sub_table = border_cell.add_table(rows=1, cols=1)
    sub_cell = sub_table.rows[0].cells[0]
    apply_cell_borders(sub_cell)  # Add borders
    run = sub_cell.paragraphs[0].add_run("✨ Source Code ✨")
    apply_font(run, size=14, bold=True, font_name="Segoe UI")
    apply_shading(sub_cell, "E6F7FF")  # Light sky blue

    # Code block with light yellow background
    code_table = border_cell.add_table(rows=1, cols=1)
    code_cell = code_table.rows[0].cells[0]
    apply_cell_borders(code_cell)  # Add borders
    apply_shading(code_cell, "FFF9E6")  # Light yellow

    if highlight:
        add_syntax_highlighted_code_to_cell(code_cell, content, file_path, font_size=11)
    else:
        run = code_cell.paragraphs[0].add_run(content)
        apply_font(run, size=11, font_name="Consolas")

    # AI Output Section (NEW - only if ai_output is provided)
    if ai_output:
        # Add spacing between source and output
        border_cell.add_paragraph()
        
        # Output subheading with emoji and creative styling
        output_sub_table = border_cell.add_table(rows=1, cols=1)
        output_sub_cell = output_sub_table.rows[0].cells[0]
        apply_cell_borders(output_sub_cell)
        run = output_sub_cell.paragraphs[0].add_run("🚀 Output 🚀")
        apply_font(run, size=14, bold=True, font_name="Segoe UI")
        apply_shading(output_sub_cell, "E6FFE6")  # Light green

        # Output content block with light mint background
        output_content_table = border_cell.add_table(rows=1, cols=1)
        output_content_cell = output_content_table.rows[0].cells[0]
        apply_cell_borders(output_content_cell)
        apply_shading(output_content_cell, "F0FFF0")  # Light mint green

        # Add the AI output with creative formatting
        run = output_content_cell.paragraphs[0].add_run(ai_output)
        apply_font(run, size=11, font_name="Consolas")

    # Add extra padding/spacing at bottom
    border_cell.add_paragraph()
    border_cell.add_paragraph()


# -------------------------------
# Updated Main Generator
# -------------------------------

def generate_docx(files: list, template: str, output_path: str, syntax_highlight=False, 
                  include_credentials=False, credentials=None, 
                  index_auto_generation=False, index_fields=None,
                  page_numbering=False, ai_outputs=None):
    """
    Generate a DOCX document using a chosen template with optional global features.
    - files: list of file paths to include
    - template: template1 | template2 | template3 | template4 | template5
    - output_path: output .docx path
    - syntax_highlight: True = syntax-colored code, False = plain text
    - include_credentials: True = add student details section
    - credentials: dict of credential fields and values
    - index_auto_generation: True = add index page
    - index_fields: dict/JSON of boolean fields for index table
    - page_numbering: True = add page numbers to footer
    """
    doc = Document()

    # Set narrow margins globally
    set_doc_margins(doc, margin_inch=0.5)

    # Add credentials section if enabled
    if include_credentials and credentials:
        add_credentials(doc, credentials)
        doc.add_page_break()

    # Add index page if enabled
    if index_auto_generation:
        add_index_page(doc, index_fields, len(files))

    # Process files with selected template
    for idx, file_path in enumerate(files, 1):
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        # Get AI output if available
        ai_output = ai_outputs.get(file_path) if ai_outputs else None
    
        if template == "template1":
            apply_template1(doc, idx, content, file_path, syntax_highlight, ai_output)
        elif template == "template2":
            apply_template2(doc, idx, content, file_path, syntax_highlight, ai_output)
        elif template == "template3":
            apply_template3(doc, idx, content, file_path, syntax_highlight, ai_output)
        elif template == "template4":
            apply_template4(doc, idx, content, file_path, syntax_highlight, ai_output)
        elif template == "template5":
            apply_template5(doc, idx, content, file_path, syntax_highlight, ai_output)
        else:
            # Fallback → template1
            apply_template1(doc, idx, content, file_path, syntax_highlight, ai_output)

    # Add page numbering if enabled (must be done after all content is added)
    if page_numbering:
        add_page_numbering(doc)

    # Save document
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)